import json
import logging
import os
import re
import sqlite3
from datetime import datetime
from functools import wraps
from io import BytesIO

from dotenv import load_dotenv
from flask import (
    Flask,
    flash,
    g,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from openai import OpenAI
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.security import check_password_hash, generate_password_hash

load_dotenv()

DATABASE = os.getenv("DATABASE_PATH", "users.db")


def env_flag(name, default=False):
    """
    Read an environment variable and convert common truthy strings to True.
    """
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


FLASK_ENV = os.getenv("FLASK_ENV", "development").strip().lower()
FLASK_DEBUG = env_flag("FLASK_DEBUG", default=(FLASK_ENV == "development"))

# Session keys used to preserve page state when the user navigates away and comes back.
INDEX_DRAFT_SESSION_KEY = "index_draft_v1"
INTERVIEW_DRAFT_SESSION_KEY = "interview_draft_v1"

app = Flask(__name__)

app.secret_key = os.getenv("FLASK_SECRET_KEY")
if not app.secret_key:
    raise RuntimeError("FLASK_SECRET_KEY is not set. Add it to your .env file and restart the app.")

app.config["DATABASE"] = os.getenv("DATABASE_PATH", "users.db")
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB

openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key) if openai_api_key else None

log_level = logging.DEBUG if FLASK_DEBUG else logging.INFO
logging.basicConfig(level=log_level)
app.logger.setLevel(log_level)


# -----------------------------
# Database Helpers
# -----------------------------
def get_db():
    """
    Open one SQLite connection per request and store it on Flask's g object.
    """
    if "db" not in g:
        g.db = sqlite3.connect(app.config["DATABASE"])
        g.db.row_factory = sqlite3.Row
    return g.db


@app.teardown_appcontext
def close_db(exception=None):
    """
    Close the database connection at the end of the request.
    """
    db = g.pop("db", None)
    if db is not None:
        db.close()


def get_table_columns(table_name):
    """
    Return a set of column names for a table.

    This is used for safe, additive migration logic.
    """
    db = get_db()
    rows = db.execute(f"PRAGMA table_info({table_name})").fetchall()
    return {row["name"] for row in rows}


def ensure_column(table_name, column_name, column_definition):
    """
    Add a column only if it does not already exist.

    Important:
    CREATE TABLE IF NOT EXISTS does not update old tables. This helper allows
    us to safely add missing columns without deleting or recreating tables.
    """
    columns = get_table_columns(table_name)
    if column_name in columns:
        return

    db = get_db()
    db.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_definition}")
    db.commit()


def now_timestamp():
    """
    Return a simple timestamp string.
    """
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def ensure_interview_sessions_schema():
    """
    Normalize interview_sessions across older and newer app versions.

    Important compatibility cases we now support:
    - older DBs may have: questions
    - other older DBs may have: interview_questions
    - newer DBs expect: questions_text

    We do NOT remove old columns.
    We only add the new column and copy old data into it when needed.

    This keeps existing user data safe.
    """
    columns = get_table_columns("interview_sessions")
    db = get_db()

    # Newer code expects questions_text.
    if "questions_text" not in columns:
        db.execute("ALTER TABLE interview_sessions ADD COLUMN questions_text TEXT NOT NULL DEFAULT ''")
        db.commit()
        columns = get_table_columns("interview_sessions")

    # If an older column named `questions` exists, backfill into questions_text.
    if "questions" in columns:
        db.execute(
            "UPDATE interview_sessions "
            "SET questions_text = questions "
            "WHERE questions_text = '' OR questions_text IS NULL"
        )
        db.commit()

    # If an older column named `interview_questions` exists, backfill into questions_text.
    if "interview_questions" in columns:
        db.execute(
            "UPDATE interview_sessions "
            "SET questions_text = interview_questions "
            "WHERE questions_text = '' OR questions_text IS NULL"
        )
        db.commit()

    # Add created_at to older DBs that do not have it yet.
    columns = get_table_columns("interview_sessions")
    if "created_at" not in columns:
        db.execute("ALTER TABLE interview_sessions ADD COLUMN created_at TEXT NOT NULL DEFAULT ''")
        db.execute(
            "UPDATE interview_sessions "
            "SET created_at = ? "
            "WHERE created_at = '' OR created_at IS NULL",
            (now_timestamp(),),
        )
        db.commit()


def get_interview_questions_column_name():
    """
    Return whichever interview-questions column exists in the current database.

    Priority:
    1. questions_text (new)
    2. interview_questions (older version that caused your NOT NULL error)
    3. questions (older version)
    4. fallback to questions_text
    """
    columns = get_table_columns("interview_sessions")

    if "questions_text" in columns:
        return "questions_text"
    if "interview_questions" in columns:
        return "interview_questions"
    if "questions" in columns:
        return "questions"

    return "questions_text"


def init_db():
    """
    Initialize the database and run safe additive migrations.

    This preserves existing user data and older tables.
    """
    db = get_db()

    db.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL
        )
        """
    )

    db.execute(
        """
        CREATE TABLE IF NOT EXISTS applications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            resume_text TEXT NOT NULL,
            job_posting TEXT NOT NULL,
            tailored_bullets TEXT NOT NULL,
            cover_letter TEXT NOT NULL,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
        """
    )

    db.execute(
        """
        CREATE TABLE IF NOT EXISTS interview_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            resume_text TEXT NOT NULL,
            job_posting TEXT NOT NULL,
            questions_text TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
        """
    )

    db.execute(
        """
        CREATE TABLE IF NOT EXISTS interview_responses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER NOT NULL,
            question_text TEXT NOT NULL,
            generated_answer TEXT NOT NULL DEFAULT '',
            user_answer_text TEXT NOT NULL DEFAULT '',
            feedback_output TEXT NOT NULL DEFAULT '',
            updated_at TEXT NOT NULL,
            FOREIGN KEY (session_id) REFERENCES interview_sessions (id)
        )
        """
    )

    db.commit()

    # Safe additive fixes for older DB versions.
    ensure_column("interview_responses", "generated_answer", "TEXT NOT NULL DEFAULT ''")
    ensure_column("interview_responses", "user_answer_text", "TEXT NOT NULL DEFAULT ''")
    ensure_column("interview_responses", "feedback_output", "TEXT NOT NULL DEFAULT ''")
    ensure_column("interview_responses", "updated_at", "TEXT NOT NULL DEFAULT ''")
    ensure_interview_sessions_schema()


# -----------------------------
# Auth Helpers
# -----------------------------
def login_required(view_func):
    """
    Require the user to be logged in before accessing the wrapped route.
    """
    @wraps(view_func)
    def wrapped_view(*args, **kwargs):
        if "user_id" not in session:
            flash("Please log in first.", "warning")
            return redirect(url_for("login"))
        return view_func(*args, **kwargs)

    return wrapped_view


# -----------------------------
# PDF Extraction
# -----------------------------
def extract_text_from_pdf(file_storage):
    """
    Extract plain text from an uploaded PDF.
    """
    if not file_storage or not file_storage.filename:
        raise ValueError("No PDF file was uploaded.")

    filename = file_storage.filename.lower().strip()
    if not filename.endswith(".pdf"):
        raise ValueError("Only PDF files are supported.")

    try:
        pdf_bytes = file_storage.read()
        if not pdf_bytes:
            raise ValueError("The uploaded PDF appears to be empty.")

        reader = PdfReader(BytesIO(pdf_bytes))

        if getattr(reader, "is_encrypted", False):
            raise ValueError(
                "This PDF is encrypted or password-protected and could not be read."
            )

        extracted_pages = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_pages.append(page_text.strip())

        full_text = "\n\n".join(extracted_pages).strip()

        if not full_text:
            raise ValueError(
                "No readable text was found in this PDF. It may be image-based or scanned."
            )

        return full_text

    except ValueError:
        raise
    except Exception as exc:
        app.logger.exception("PDF extraction failed.")
        raise ValueError(f"Could not read the PDF: {exc}") from exc


# -----------------------------
# Text / Formatting Helpers
# -----------------------------
def sanitize_model_text(text):
    """
    Normalize possibly-empty text values from forms, DB rows, or AI output.
    """
    return (text or "").strip()


def extract_tagged_section(content, start_tag, end_tag):
    """
    Pull a tagged block out of the model response.
    """
    start_index = content.find(start_tag)
    end_index = content.find(end_tag)

    if start_index == -1 or end_index == -1 or end_index <= start_index:
        return ""

    start_index += len(start_tag)
    return content[start_index:end_index].strip()


def parse_job_target_info(text):
    """
    Parse the detected job title and company block returned by the AI.
    """
    cleaned = sanitize_model_text(text)
    result = {
        "job_title": "Not detected",
        "company_name": "Not detected",
    }

    if not cleaned:
        return result

    for raw_line in cleaned.splitlines():
        line = raw_line.strip()
        lower_line = line.lower()

        if lower_line.startswith("job title:"):
            value = line.split(":", 1)[1].strip()
            if value:
                result["job_title"] = value

        elif lower_line.startswith("company:"):
            value = line.split(":", 1)[1].strip()
            if value:
                result["company_name"] = value

    return result


def format_tailored_bullets(text):
    """
    Normalize the tailored-bullets block for cleaner display.
    """
    cleaned = sanitize_model_text(text)
    if not cleaned:
        return ""

    lines = [line.rstrip() for line in cleaned.splitlines()]
    formatted_lines = []
    previous_blank = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if not previous_blank:
                formatted_lines.append("")
            previous_blank = True
            continue

        previous_blank = False

        if stripped.upper().startswith("MATCH SUMMARY"):
            formatted_lines.append("MATCH SUMMARY")
        elif stripped.upper().startswith("TAILORED BULLETS"):
            formatted_lines.append("TAILORED BULLETS")
        elif stripped.upper().startswith("KEYWORDS TO MIRROR"):
            formatted_lines.append("KEYWORDS TO MIRROR")
        else:
            formatted_lines.append(stripped)

    return "\n".join(formatted_lines).strip()


def format_cover_letter(text):
    """
    Normalize the cover-letter block into readable paragraphs.
    """
    cleaned = sanitize_model_text(text)
    if not cleaned:
        return ""

    lines = [line.strip() for line in cleaned.splitlines()]
    paragraphs = []
    current_paragraph = []

    for line in lines:
        if not line:
            if current_paragraph:
                paragraphs.append(" ".join(current_paragraph).strip())
                current_paragraph = []
            continue
        current_paragraph.append(line)

    if current_paragraph:
        paragraphs.append(" ".join(current_paragraph).strip())

    return "\n\n".join(paragraphs).strip()


def parse_interview_questions_text(questions_text):
    """
    Convert a stored interview question block back into a list of questions.

    Important:
    The Interview Coach stores the raw question block as text. This helper is
    used to split it back into display-ready lines for interview history and
    saved-session reloads.
    """
    cleaned = sanitize_model_text(questions_text)
    if not cleaned:
        return []
    return [line.strip() for line in cleaned.splitlines() if line.strip()]


# -----------------------------
# Interview Coach AI Helpers
# -----------------------------
def generate_interview_questions(resume_text, job_posting):
    """
    Generate interview questions tailored to the user's resume and job posting.
    """
    if not client:
        raise RuntimeError(
            "OPENAI_API_KEY is missing. Add it to your .env file and restart the app."
        )

    prompt = f"""
You are an expert hiring manager and interview coach.

Given this resume:
{resume_text}

And this job description:
{job_posting}

Generate exactly 7 interview questions tailored to this candidate and this role.

Requirements:
- Mix behavioral and role-specific questions
- Make the questions realistic and professional
- Focus on what the interviewer would most likely ask this candidate
- Return the output as a plain numbered list only
- Do not include extra intro text
""".strip()

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a professional interview coach."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.7,
        )
    except Exception as exc:
        app.logger.exception("Interview question generation failed.")
        raise RuntimeError(f"Interview question generation failed: {exc}") from exc

    choices = getattr(response, "choices", None)
    if not choices:
        raise RuntimeError("The AI returned no interview questions. Please try again.")

    first_message = getattr(choices[0], "message", None)
    content = sanitize_model_text(getattr(first_message, "content", ""))

    if not content:
        raise RuntimeError("The AI returned an empty interview questions response.")

    return content


def generate_interview_answer(question, resume_text, job_posting):
    """
    Generate a sample interview answer.
    """
    if not client:
        raise RuntimeError(
            "OPENAI_API_KEY is missing. Add it to your .env file and restart the app."
        )

    prompt = f"""
You are an expert interview coach.

Given this resume:
{resume_text}

And this job description:
{job_posting}

Answer this interview question:
{question}

Requirements:
- Use the STAR method (Situation, Task, Action, Result) when appropriate
- Be concise but impactful (150-250 words)
- Sound natural and confident
- Do not invent fake experience
- Keep the answer truthful to the resume
""".strip()

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a professional interview coach."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.7,
        )
    except Exception as exc:
        app.logger.exception("Interview answer generation failed.")
        raise RuntimeError(f"Interview answer generation failed: {exc}") from exc

    choices = getattr(response, "choices", None)
    if not choices:
        raise RuntimeError("The AI returned no interview answer. Please try again.")

    first_message = getattr(choices[0], "message", None)
    content = sanitize_model_text(getattr(first_message, "content", ""))

    if not content:
        raise RuntimeError("The AI returned an empty interview answer.")

    return content


def generate_interview_feedback(question, user_answer, resume_text, job_posting):
    """
    Generate structured feedback for a user's interview answer.
    """
    if not client:
        raise RuntimeError(
            "OPENAI_API_KEY is missing. Add it to your .env file and restart the app."
        )

    prompt = f"""
You are an expert interview coach.

Evaluate the candidate's answer to this interview question.

QUESTION:
{question}

CANDIDATE ANSWER:
{user_answer}

CONTEXT:
Resume:
{resume_text}

Job Description:
{job_posting}

Provide structured feedback in this exact format:

Score:
[Give a score from 1-10]

Strengths:
- [strength]
- [strength]

Areas for Improvement:
- [improvement]
- [improvement]

Improved Answer:
[Write a stronger version of the answer using the STAR method when appropriate]

Requirements:
- Be honest, constructive, and practical
- Keep the improved answer truthful to the resume
- Do not invent fake experience
""".strip()

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a professional interview coach."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.7,
        )
    except Exception as exc:
        app.logger.exception("Interview feedback generation failed.")
        raise RuntimeError(f"Interview feedback generation failed: {exc}") from exc

    choices = getattr(response, "choices", None)
    if not choices:
        raise RuntimeError("The AI returned no feedback. Please try again.")

    first_message = getattr(choices[0], "message", None)
    content = sanitize_model_text(getattr(first_message, "content", ""))

    if not content:
        raise RuntimeError("The AI returned an empty feedback response.")

    return content


# -----------------------------
# Interview Session / History Helpers
# -----------------------------
def create_interview_session(user_id, resume_text, job_posting, questions_text):
    """
    Save a new interview session.

    Critical compatibility logic:
    Older databases may still require the NOT NULL column `interview_questions`.
    Newer code uses `questions_text`.

    So this function detects the real schema and writes to all relevant question
    columns that exist.
    """
    db = get_db()
    columns = get_table_columns("interview_sessions")

    cleaned_resume = sanitize_model_text(resume_text)
    cleaned_job_posting = sanitize_model_text(job_posting)
    cleaned_questions = sanitize_model_text(questions_text)
    created_at_value = now_timestamp()

    insert_columns = ["user_id", "resume_text", "job_posting"]
    insert_values = [user_id, cleaned_resume, cleaned_job_posting]

    # Newer schema
    if "questions_text" in columns:
        insert_columns.append("questions_text")
        insert_values.append(cleaned_questions)

    # Older schema variant that caused your IntegrityError
    if "interview_questions" in columns:
        insert_columns.append("interview_questions")
        insert_values.append(cleaned_questions)

    # Another older schema variant
    if "questions" in columns:
        insert_columns.append("questions")
        insert_values.append(cleaned_questions)

    if "created_at" in columns:
        insert_columns.append("created_at")
        insert_values.append(created_at_value)

    placeholders = ", ".join(["?"] * len(insert_columns))
    columns_sql = ", ".join(insert_columns)

    cursor = db.execute(
        f"""
        INSERT INTO interview_sessions ({columns_sql})
        VALUES ({placeholders})
        """,
        tuple(insert_values),
    )
    db.commit()
    return cursor.lastrowid


def get_interview_session_for_user(user_id, session_id):
    """
    Load one interview session for the current user.

    We alias the detected question column back to `questions_text` so the rest of
    the route/template code can stay consistent.
    """
    if not session_id:
        return None

    db = get_db()
    questions_column = get_interview_questions_column_name()

    return db.execute(
        f"""
        SELECT id, user_id, resume_text, job_posting, {questions_column} AS questions_text, created_at
        FROM interview_sessions
        WHERE id = ? AND user_id = ?
        """,
        (session_id, user_id),
    ).fetchone()


def ensure_interview_session_id(current_session_id, user_id, resume_text, job_posting, questions_text):
    """
    Return a valid interview session ID.

    If the posted session ID is missing or invalid, create a new one.
    """
    try:
        session_id = int(str(current_session_id).strip()) if str(current_session_id).strip() else None
    except (TypeError, ValueError):
        session_id = None

    existing = get_interview_session_for_user(user_id, session_id) if session_id else None
    if existing:
        return existing["id"]

    return create_interview_session(user_id, resume_text, job_posting, questions_text)


def save_interview_response_record(
    session_id,
    question_text,
    generated_answer="",
    user_answer_text="",
    feedback_output="",
):
    """
    Save one interview attempt row.

    Each row may contain:
    - a generated sample answer
    - a user-written answer
    - AI feedback
    """
    if not session_id or not sanitize_model_text(question_text):
        return

    db = get_db()
    db.execute(
        """
        INSERT INTO interview_responses (
            session_id, question_text, generated_answer, user_answer_text, feedback_output, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            session_id,
            sanitize_model_text(question_text),
            sanitize_model_text(generated_answer),
            sanitize_model_text(user_answer_text),
            sanitize_model_text(feedback_output),
            now_timestamp(),
        ),
    )
    db.commit()


def fetch_application_history_for_user(user_id):
    """
    Load application history only.

    This keeps application history clearly separated from interview history.
    """
    db = get_db()
    return db.execute(
        """
        SELECT id, resume_text, job_posting, tailored_bullets, cover_letter
        FROM applications
        WHERE user_id = ?
        ORDER BY id DESC
        """,
        (user_id,),
    ).fetchall()


def build_interview_history_sessions(user_id):
    """
    Build the grouped interview history structure used by interview_history.html.

    This is intentionally separate from the application history flow.
    """
    db = get_db()
    questions_column = get_interview_questions_column_name()

    session_rows = db.execute(
        f"""
        SELECT id, resume_text, job_posting, {questions_column} AS questions_text, created_at
        FROM interview_sessions
        WHERE user_id = ?
        ORDER BY id DESC
        """,
        (user_id,),
    ).fetchall()

    sessions = []

    for row in session_rows:
        response_rows = db.execute(
            """
            SELECT id, question_text, generated_answer, user_answer_text, feedback_output, updated_at
            FROM interview_responses
            WHERE session_id = ?
            ORDER BY id ASC
            """,
            (row["id"],),
        ).fetchall()

        grouped_attempts = {}

        for response in response_rows:
            question_text = sanitize_model_text(response["question_text"])
            if not question_text:
                continue

            if question_text not in grouped_attempts:
                grouped_attempts[question_text] = []

            grouped_attempts[question_text].append({
                "generated_answer": sanitize_model_text(response["generated_answer"]),
                "user_answer_text": sanitize_model_text(response["user_answer_text"]),
                "feedback_output": sanitize_model_text(response["feedback_output"]),
                "updated_at": sanitize_model_text(response["updated_at"]) or "Not available",
            })

        sessions.append({
            "id": row["id"],
            "resume_text": sanitize_model_text(row["resume_text"]),
            "job_posting": sanitize_model_text(row["job_posting"]),
            "questions": parse_interview_questions_text(row["questions_text"]),
            "responses": grouped_attempts,
            "response_count": sum(len(attempts) for attempts in grouped_attempts.values()),
            "created_at": sanitize_model_text(row["created_at"]) or "Not available",
        })

    return sessions


# -----------------------------
# Version Helpers
# -----------------------------
def parse_previous_versions(raw_json):
    """
    Parse the hidden JSON used for Home-page version history.
    """
    if not raw_json:
        return []

    try:
        parsed = json.loads(raw_json)
    except (json.JSONDecodeError, TypeError):
        return []

    if not isinstance(parsed, list):
        return []

    cleaned_versions = []
    for item in parsed:
        if not isinstance(item, dict):
            continue

        cleaned_versions.append(
            {
                "job_title": sanitize_model_text(item.get("job_title", "")),
                "company_name": sanitize_model_text(item.get("company_name", "")),
                "tailored_bullets": sanitize_model_text(item.get("tailored_bullets", "")),
                "cover_letter": sanitize_model_text(item.get("cover_letter", "")),
                "note": sanitize_model_text(item.get("note", "")),
            }
        )

    return cleaned_versions


def build_result_snapshot(context):
    """
    Capture the current generated result for Home-page versioning.
    """
    if not context.get("tailored_bullets") and not context.get("cover_letter"):
        return None

    return {
        "job_title": sanitize_model_text(context.get("job_title", "")),
        "company_name": sanitize_model_text(context.get("company_name", "")),
        "tailored_bullets": sanitize_model_text(context.get("tailored_bullets", "")),
        "cover_letter": sanitize_model_text(context.get("cover_letter", "")),
        "note": "",
    }


def snapshot_exists(snapshot, previous_versions):
    """
    Prevent duplicate saved snapshots in the Home-page version list.
    """
    if not snapshot:
        return False

    comparable_snapshot = {
        "job_title": snapshot.get("job_title", ""),
        "company_name": snapshot.get("company_name", ""),
        "tailored_bullets": snapshot.get("tailored_bullets", ""),
        "cover_letter": snapshot.get("cover_letter", ""),
    }

    for version in previous_versions:
        comparable_version = {
            "job_title": version.get("job_title", ""),
            "company_name": version.get("company_name", ""),
            "tailored_bullets": version.get("tailored_bullets", ""),
            "cover_letter": version.get("cover_letter", ""),
        }
        if comparable_version == comparable_snapshot:
            return True

    return False


def normalize_version_item(item):
    """
    Normalize one previous-version item.
    """
    if not isinstance(item, dict):
        return None

    return {
        "job_title": sanitize_model_text(item.get("job_title", "")),
        "company_name": sanitize_model_text(item.get("company_name", "")),
        "tailored_bullets": sanitize_model_text(item.get("tailored_bullets", "")),
        "cover_letter": sanitize_model_text(item.get("cover_letter", "")),
        "note": sanitize_model_text(item.get("note", "")),
    }


def pop_previous_version(previous_versions, index):
    """
    Remove one previous version and return it so it can become current.
    """
    if index < 0 or index >= len(previous_versions):
        return None, previous_versions

    selected = normalize_version_item(previous_versions[index])
    remaining = previous_versions[:index] + previous_versions[index + 1:]
    return selected, remaining


def remove_previous_version(previous_versions, index):
    """
    Delete one previous-version entry.
    """
    if index < 0 or index >= len(previous_versions):
        return None, previous_versions

    removed = normalize_version_item(previous_versions[index])
    remaining = previous_versions[:index] + previous_versions[index + 1:]
    return removed, remaining


def update_previous_version_note(previous_versions, index, note):
    """
    Save the note/label for one previous version.
    """
    if index < 0 or index >= len(previous_versions):
        return False, previous_versions

    updated_versions = []
    for i, version in enumerate(previous_versions):
        normalized = normalize_version_item(version)
        if not normalized:
            continue

        if i == index:
            normalized["note"] = sanitize_model_text(note)

        updated_versions.append(normalized)

    return True, updated_versions


def get_previous_version_by_index(previous_versions, index):
    """
    Return one previous version by index.
    """
    if index < 0 or index >= len(previous_versions):
        return None
    return normalize_version_item(previous_versions[index])


# -----------------------------
# Session Draft Helpers - Home Page
# -----------------------------
def empty_index_context():
    """
    Default Home-page state.
    """
    return {
        "resume_text": "",
        "job_posting": "",
        "job_title": "",
        "company_name": "",
        "tailored_bullets": "",
        "cover_letter": "",
        "extracted_resume_preview": "",
        "previous_versions": [],
    }


def normalize_index_context(context):
    """
    Sanitize Home-page state before storing or rendering it.
    """
    base = empty_index_context()

    if not isinstance(context, dict):
        return base

    base["resume_text"] = sanitize_model_text(context.get("resume_text", ""))
    base["job_posting"] = sanitize_model_text(context.get("job_posting", ""))
    base["job_title"] = sanitize_model_text(context.get("job_title", ""))
    base["company_name"] = sanitize_model_text(context.get("company_name", ""))
    base["tailored_bullets"] = sanitize_model_text(context.get("tailored_bullets", ""))
    base["cover_letter"] = sanitize_model_text(context.get("cover_letter", ""))
    base["extracted_resume_preview"] = sanitize_model_text(
        context.get("extracted_resume_preview", "")
    )

    raw_versions = context.get("previous_versions", [])
    cleaned_versions = []
    if isinstance(raw_versions, list):
        for item in raw_versions:
            normalized = normalize_version_item(item)
            if normalized:
                cleaned_versions.append(normalized)

    base["previous_versions"] = cleaned_versions
    return base


def save_index_draft(context):
    """
    Persist Home-page draft state in the Flask session.
    """
    session[INDEX_DRAFT_SESSION_KEY] = normalize_index_context(context)


def load_index_draft():
    """
    Reload Home-page draft state from the Flask session.
    """
    saved = session.get(INDEX_DRAFT_SESSION_KEY)
    if not saved:
        return empty_index_context()
    return normalize_index_context(saved)


# -----------------------------
# Session Draft Helpers - Interview Coach
# -----------------------------
def empty_interview_context():
    """
    Default Interview Coach state.

    This exists because Interview Coach needs the same kind of state persistence
    that your Home page already uses. Without this, leaving the page and coming
    back can clear out the current practice session UI state.
    """
    return {
        "resume_text": "",
        "job_posting": "",
        "interview_questions": "",
        "generated_answer": "",
        "selected_question": "",
        "feedback_output": "",
        "feedback_question": "",
        "user_answer_text": "",
        "all_questions": "",
        "current_interview_session_id": "",
    }


def normalize_interview_context(context):
    """
    Sanitize Interview Coach state before storing or rendering it.
    """
    base = empty_interview_context()

    if not isinstance(context, dict):
        return base

    base["resume_text"] = sanitize_model_text(context.get("resume_text", ""))
    base["job_posting"] = sanitize_model_text(context.get("job_posting", ""))
    base["interview_questions"] = sanitize_model_text(context.get("interview_questions", ""))
    base["generated_answer"] = sanitize_model_text(context.get("generated_answer", ""))
    base["selected_question"] = sanitize_model_text(context.get("selected_question", ""))
    base["feedback_output"] = sanitize_model_text(context.get("feedback_output", ""))
    base["feedback_question"] = sanitize_model_text(context.get("feedback_question", ""))
    base["user_answer_text"] = sanitize_model_text(context.get("user_answer_text", ""))
    base["all_questions"] = sanitize_model_text(context.get("all_questions", ""))
    base["current_interview_session_id"] = sanitize_model_text(
        context.get("current_interview_session_id", "")
    )

    return base


def save_interview_draft(context):
    """
    Persist Interview Coach page state in the Flask session.

    Easy-to-break area:
    This is what fixes the "toggle between pages and lose state" issue.
    """
    session[INTERVIEW_DRAFT_SESSION_KEY] = normalize_interview_context(context)


def load_interview_draft():
    """
    Reload Interview Coach page state from the Flask session.
    """
    saved = session.get(INTERVIEW_DRAFT_SESSION_KEY)
    if not saved:
        return empty_interview_context()
    return normalize_interview_context(saved)


def render_interview_with_context(context):
    """
    Normalize, persist, and render the Interview Coach state.
    """
    normalized_context = normalize_interview_context(context)
    save_interview_draft(normalized_context)
    return render_template("interview.html", **normalized_context)


# -----------------------------
# Export Helpers
# -----------------------------
def has_current_output(context):
    """
    Return True when the Home page has exportable output.
    """
    return bool(context.get("tailored_bullets") or context.get("cover_letter"))


def sanitize_filename_component(value, fallback):
    """
    Convert arbitrary text into a safe filename piece.
    """
    cleaned = sanitize_model_text(value)
    if not cleaned or cleaned.lower() in {"not detected", "not provided"}:
        cleaned = fallback

    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", cleaned).strip("_")
    return cleaned[:50] or fallback


def build_export_filename(payload, extension):
    """
    Build a stable export filename from job title and company.
    """
    job_part = sanitize_filename_component(payload.get("job_title", ""), "job")
    company_part = sanitize_filename_component(payload.get("company_name", ""), "company")
    return f"{job_part}_{company_part}_application_materials.{extension}"


def build_export_payload_from_context(context):
    """
    Convert current Home-page state into an export payload.
    """
    return {
        "job_title": sanitize_model_text(context.get("job_title", "")),
        "company_name": sanitize_model_text(context.get("company_name", "")),
        "tailored_bullets": sanitize_model_text(context.get("tailored_bullets", "")),
        "cover_letter": sanitize_model_text(context.get("cover_letter", "")),
        "note": "",
    }


def has_exportable_payload(payload):
    """
    Return True when a payload contains generated content worth exporting.
    """
    return bool(
        sanitize_model_text(payload.get("tailored_bullets", "")) or
        sanitize_model_text(payload.get("cover_letter", ""))
    )


def add_docx_multiline_block(document, text):
    """
    Add multi-line text blocks to a DOCX document while keeping rough formatting.
    """
    for line in sanitize_model_text(text).splitlines():
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)

        if stripped.startswith("-"):
            paragraph.style = "List Bullet"
            paragraph.add_run(stripped[1:].strip())
        elif stripped.isupper() and len(stripped) < 60:
            run = paragraph.add_run(stripped)
            run.bold = True
        else:
            paragraph.add_run(stripped)


def build_docx_export(payload):
    """
    Generate a DOCX file in memory and return a BytesIO object.
    """
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI Job Application Materials")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f'Exported on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    if sanitize_model_text(payload.get("note", "")):
        note_paragraph = document.add_paragraph()
        note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_paragraph.add_run(f'Label: {sanitize_model_text(payload["note"])}')
        note_run.italic = True
        note_run.font.size = Pt(10)

    document.add_paragraph("")

    meta = document.add_paragraph()
    meta.add_run("Detected Job Title: ").bold = True
    meta.add_run(payload.get("job_title") or "Not detected")

    meta2 = document.add_paragraph()
    meta2.add_run("Detected Company: ").bold = True
    meta2.add_run(payload.get("company_name") or "Not detected")

    document.add_paragraph("")

    heading = document.add_paragraph()
    heading_run = heading.add_run("Tailored Bullets")
    heading_run.bold = True
    heading_run.font.size = Pt(14)

    add_docx_multiline_block(document, payload.get("tailored_bullets", ""))

    document.add_paragraph("")

    heading2 = document.add_paragraph()
    heading2_run = heading2.add_run("Cover Letter")
    heading2_run.bold = True
    heading2_run.font.size = Pt(14)

    for paragraph_text in sanitize_model_text(payload.get("cover_letter", "")).split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if paragraph_text:
            p = document.add_paragraph(paragraph_text)
            p.paragraph_format.space_after = Pt(10)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def text_to_pdf_paragraphs(text):
    """
    Convert multi-paragraph plain text into PDF-friendly paragraph blocks.
    """
    paragraphs = []
    for block in sanitize_model_text(text).split("\n\n"):
        block = block.strip()
        if block:
            paragraphs.append(block.replace("\n", "<br/>"))
    return paragraphs


def build_pdf_export(payload):
    """
    Generate a PDF file in memory and return a BytesIO object.
    """
    output = BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#163a70"),
        alignment=1,
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "ExportSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#5b6b86"),
        alignment=1,
        spaceAfter=18,
    )
    section_style = ParagraphStyle(
        "ExportSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#163a70"),
        spaceBefore=12,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=colors.black,
        spaceAfter=8,
    )

    story = []

    story.append(Paragraph("AI Job Application Materials", title_style))
    story.append(
        Paragraph(
            f'Exported on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}',
            subtitle_style,
        )
    )

    if sanitize_model_text(payload.get("note", "")):
        story.append(
            Paragraph(
                f"<i>Label: {sanitize_model_text(payload['note'])}</i>",
                subtitle_style,
            )
        )

    story.append(
        Paragraph(
            f"<b>Detected Job Title:</b> {sanitize_model_text(payload.get('job_title')) or 'Not detected'}",
            body_style,
        )
    )
    story.append(
        Paragraph(
            f"<b>Detected Company:</b> {sanitize_model_text(payload.get('company_name')) or 'Not detected'}",
            body_style,
        )
    )

    story.append(Spacer(1, 8))
    story.append(Paragraph("Tailored Bullets", section_style))

    for paragraph_text in text_to_pdf_paragraphs(payload.get("tailored_bullets", "")):
        story.append(Paragraph(paragraph_text, body_style))

    story.append(Spacer(1, 8))
    story.append(Paragraph("Cover Letter", section_style))

    for paragraph_text in text_to_pdf_paragraphs(payload.get("cover_letter", "")):
        story.append(Paragraph(paragraph_text, body_style))

    doc.build(story)
    output.seek(0)
    return output


# -----------------------------
# OpenAI Generation
# -----------------------------
def generate_application_materials(resume_text, job_posting):
    """
    Generate tailored resume bullets and a cover letter from the main Home page.
    """
    if not client:
        raise RuntimeError(
            "OPENAI_API_KEY is missing. Add it to your .env file and restart the app."
        )

    system_message = """
You are an expert job application assistant and resume strategist.

Your job is to tailor application materials to a target role while staying fully truthful.
Never invent experience, metrics, tools, certifications, employers, titles, dates, degrees, or achievements.
You may reframe and prioritize the candidate's real experience, but you must not fabricate.
""".strip()

    user_prompt = f"""
Create three outputs from the candidate resume and the target job posting.

OUTPUT REQUIREMENTS

1) Job target section:
- Determine the most likely target job title from the posting
- Determine the company name from the posting
- If the company name is truly unavailable, use: Not provided
- Return exactly:
  Job Title: ...
  Company: ...

2) Tailored bullets section:
- Start with the heading: MATCH SUMMARY
- Under MATCH SUMMARY, write 2 concise lines summarizing why the candidate fits the role
- Then add a blank line
- Add the heading: TAILORED BULLETS
- Provide 6 bullets total
- Each bullet must:
  - begin with "-"
  - be 1-2 lines max
  - be specific and job-relevant
  - sound like polished resume content
  - emphasize transferable skills where needed
  - avoid first-person language
- Then add a blank line
- Add the heading: KEYWORDS TO MIRROR
- Provide 8-12 short comma-separated keywords or phrases pulled from the job posting that the candidate genuinely appears to match

3) Cover letter section:
- Write a professional cover letter in 3 short paragraphs
- Tone should be confident, specific, and human
- Paragraph 1: interest in the role and fit
- Paragraph 2: strongest aligned experience and value
- Paragraph 3: close professionally and express interest in next steps
- Keep it around 220-320 words
- Do not include fake facts
- Do not include placeholders like [Company Name] unless the company is truly missing from the posting
- If company name is missing, use "your team"

RETURN FORMAT
Return your response using exactly these tags and nothing else outside them:

<JOB_TARGET>
Job Title: ...
Company: ...
</JOB_TARGET>

<TAILORED_BULLETS>
...tailored bullets content here...
</TAILORED_BULLETS>

<COVER_LETTER>
...cover letter content here...
</COVER_LETTER>

CANDIDATE RESUME
{resume_text}

TARGET JOB POSTING
{job_posting}
""".strip()

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.7,
        )
    except Exception as exc:
        app.logger.exception("OpenAI request failed.")
        raise RuntimeError(f"OpenAI request failed: {exc}") from exc

    choices = getattr(response, "choices", None)
    if not choices:
        raise RuntimeError("The AI returned no choices. Please try again.")

    first_message = getattr(choices[0], "message", None)
    content = sanitize_model_text(getattr(first_message, "content", ""))

    if not content:
        raise RuntimeError("The AI returned an empty response.")

    job_target_raw = extract_tagged_section(content, "<JOB_TARGET>", "</JOB_TARGET>")
    tailored_bullets = extract_tagged_section(content, "<TAILORED_BULLETS>", "</TAILORED_BULLETS>")
    cover_letter = extract_tagged_section(content, "<COVER_LETTER>", "</COVER_LETTER>")

    if not tailored_bullets or not cover_letter:
        app.logger.warning("AI response parsing failed. Raw content: %s", content[:1000])
        raise RuntimeError(
            "The AI response could not be parsed correctly. Please try again."
        )

    job_target_info = parse_job_target_info(job_target_raw)
    tailored_bullets = format_tailored_bullets(tailored_bullets)
    cover_letter = format_cover_letter(cover_letter)

    return (
        job_target_info["job_title"],
        job_target_info["company_name"],
        tailored_bullets,
        cover_letter,
    )


# -----------------------------
# Form / Render Helpers
# -----------------------------
def normalize_form_data():
    """
    Normalize the Home page form data.
    """
    return {
        "resume_text": request.form.get("resume_text", "").strip(),
        "job_posting": request.form.get("job_posting", "").strip(),
        "job_title": request.form.get("current_job_title", "").strip(),
        "company_name": request.form.get("current_company_name", "").strip(),
        "tailored_bullets": request.form.get("current_tailored_bullets", "").strip(),
        "cover_letter": request.form.get("current_cover_letter", "").strip(),
        "extracted_resume_preview": request.form.get("extracted_resume_preview", "").strip(),
        "previous_versions": parse_previous_versions(
            request.form.get("previous_versions_json", "")
        ),
    }


def render_index_with_context(context):
    """
    Normalize, persist, and render the Home page state.
    """
    normalized_context = normalize_index_context(context)
    save_index_draft(normalized_context)
    return render_template("index.html", **normalized_context)


def parse_selected_version_index():
    """
    Parse the hidden selected previous-version index from the form.
    """
    selected_index_raw = request.form.get("selected_version_index", "").strip()

    try:
        return int(selected_index_raw)
    except (TypeError, ValueError):
        return None


# -----------------------------
# Error Handlers
# -----------------------------
@app.errorhandler(RequestEntityTooLarge)
def handle_large_file(error):
    """
    Friendly upload-size error.
    """
    flash("The uploaded file is too large. Please upload a PDF under 10 MB.", "error")
    if "user_id" in session:
        return redirect(url_for("index"))
    return redirect(url_for("login"))


@app.errorhandler(500)
def handle_internal_server_error(error):
    """
    Catch-all server error handler.
    """
    app.logger.exception("Unhandled server error: %s", error)
    flash("Something went wrong on the server. Please try again.", "error")
    if "user_id" in session:
        return redirect(url_for("index"))
    return redirect(url_for("login"))


# -----------------------------
# Routes
# -----------------------------
@app.route("/")
@login_required
def root():
    """
    Redirect root to the main Home page.
    """
    return redirect(url_for("index"))


@app.route("/register", methods=["GET", "POST"])
def register():
    """
    Register a new user account.
    """
    init_db()

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        if not username or not password:
            flash("Username and password are required.", "error")
            return render_template("register.html")

        db = get_db()
        existing_user = db.execute(
            "SELECT id FROM users WHERE username = ?",
            (username,),
        ).fetchone()

        if existing_user:
            flash("That username is already taken.", "error")
            return render_template("register.html")

        password_hash = generate_password_hash(password)
        db.execute(
            "INSERT INTO users (username, password_hash) VALUES (?, ?)",
            (username, password_hash),
        )
        db.commit()

        flash("Registration successful. Please log in.", "success")
        return redirect(url_for("login"))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    """
    Log an existing user into the app.
    """
    init_db()

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        if not username or not password:
            flash("Username and password are required.", "error")
            return render_template("login.html")

        db = get_db()
        user = db.execute(
            "SELECT * FROM users WHERE username = ?",
            (username,),
        ).fetchone()

        if user and check_password_hash(user["password_hash"], password):
            session.clear()
            session["user_id"] = user["id"]
            session["username"] = user["username"]
            flash("Logged in successfully.", "success")
            return redirect(url_for("index"))

        flash("Invalid username or password.", "error")

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    """
    Log out the current user.

    Because Flask session storage is cleared here, both the Home page draft and
    the Interview Coach draft are cleared too.
    """
    session.clear()
    flash("You have been logged out.", "success")
    return redirect(url_for("login"))


@app.route("/index", methods=["GET", "POST"])
@login_required
def index():
    """
    Main job application assistant page.

    This route is separate from:
    - /history
    - /interview
    - /interview_history
    """
    init_db()

    context = load_index_draft()

    if request.method == "POST":
        context = normalize_form_data()
        action = request.form.get("action", "generate").strip()

        if action == "extract_pdf":
            uploaded_resume = request.files.get("resume_pdf")

            try:
                extracted_text = extract_text_from_pdf(uploaded_resume)
                context["resume_text"] = extracted_text
                context["extracted_resume_preview"] = extracted_text
                flash(
                    "PDF uploaded and text extracted successfully. Review it below before generating.",
                    "success",
                )
            except ValueError as exc:
                flash(str(exc), "error")

            return render_index_with_context(context)

        if action == "export_docx":
            payload = build_export_payload_from_context(context)

            if not has_exportable_payload(payload):
                flash("There is no current version to export.", "error")
                return render_index_with_context(context)

            try:
                save_index_draft(context)
                file_data = build_docx_export(payload)
                filename = build_export_filename(payload, "docx")
                return send_file(
                    file_data,
                    as_attachment=True,
                    download_name=filename,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception:
                app.logger.exception("DOCX export failed.")
                flash("DOCX export failed.", "error")
                return render_index_with_context(context)

        if action == "export_pdf":
            payload = build_export_payload_from_context(context)

            if not has_exportable_payload(payload):
                flash("There is no current version to export.", "error")
                return render_index_with_context(context)

            try:
                save_index_draft(context)
                file_data = build_pdf_export(payload)
                filename = build_export_filename(payload, "pdf")
                return send_file(
                    file_data,
                    as_attachment=True,
                    download_name=filename,
                    mimetype="application/pdf",
                )
            except Exception:
                app.logger.exception("PDF export failed.")
                flash("PDF export failed.", "error")
                return render_index_with_context(context)

        if action in ("export_selected_docx", "export_selected_pdf"):
            selected_index = parse_selected_version_index()

            if selected_index is None:
                flash("That previous version could not be exported.", "error")
                return render_index_with_context(context)

            payload = get_previous_version_by_index(context["previous_versions"], selected_index)

            if not payload or not has_exportable_payload(payload):
                flash("That previous version could not be exported.", "error")
                return render_index_with_context(context)

            try:
                save_index_draft(context)

                if action == "export_selected_docx":
                    file_data = build_docx_export(payload)
                    filename = build_export_filename(payload, "docx")
                    mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    file_data = build_pdf_export(payload)
                    filename = build_export_filename(payload, "pdf")
                    mimetype = "application/pdf"

                return send_file(
                    file_data,
                    as_attachment=True,
                    download_name=filename,
                    mimetype=mimetype,
                )
            except Exception:
                app.logger.exception("Selected previous version export failed.")
                flash("Selected previous version export failed.", "error")
                return render_index_with_context(context)

        if action == "clear_versions":
            if context["previous_versions"]:
                context["previous_versions"] = []
                flash("Previous versions cleared from the page.", "success")
            else:
                flash("There were no previous versions to clear.", "warning")

            return render_index_with_context(context)

        if action == "save_version_note":
            selected_index = parse_selected_version_index()
            version_note = request.form.get("version_note", "").strip()

            if selected_index is None:
                flash("That version note could not be saved.", "error")
                return render_index_with_context(context)

            updated, updated_versions = update_previous_version_note(
                context["previous_versions"],
                selected_index,
                version_note,
            )

            if not updated:
                flash("That version note could not be saved.", "error")
                return render_index_with_context(context)

            context["previous_versions"] = updated_versions
            flash("Version note saved.", "success")
            return render_index_with_context(context)

        if action == "delete_version":
            selected_index = parse_selected_version_index()

            if selected_index is None:
                flash("That version could not be deleted.", "error")
                return render_index_with_context(context)

            removed_version, remaining_versions = remove_previous_version(
                context["previous_versions"],
                selected_index,
            )

            if not removed_version:
                flash("That version could not be deleted.", "error")
                return render_index_with_context(context)

            context["previous_versions"] = remaining_versions
            flash("Selected previous version was deleted.", "success")
            return render_index_with_context(context)

        if action == "use_version":
            selected_index = parse_selected_version_index()

            if selected_index is None:
                flash("That version could not be selected.", "error")
                return render_index_with_context(context)

            selected_version, remaining_versions = pop_previous_version(
                context["previous_versions"],
                selected_index,
            )

            if not selected_version:
                flash("That version could not be selected.", "error")
                return render_index_with_context(context)

            current_snapshot = build_result_snapshot(context)

            if current_snapshot and not snapshot_exists(current_snapshot, remaining_versions):
                remaining_versions.insert(0, current_snapshot)

            context["previous_versions"] = remaining_versions
            context["job_title"] = selected_version["job_title"]
            context["company_name"] = selected_version["company_name"]
            context["tailored_bullets"] = selected_version["tailored_bullets"]
            context["cover_letter"] = selected_version["cover_letter"]

            flash("Selected previous version is now the current version.", "success")
            return render_index_with_context(context)

        if action in ("generate", "regenerate"):
            resume_text = context["resume_text"]
            job_posting = context["job_posting"]

            if not resume_text:
                flash(
                    "Please paste your resume text or upload a PDF resume first.",
                    "error",
                )
                return render_index_with_context(context)

            if len(resume_text) < 50:
                flash(
                    "Resume text looks too short. Please paste more complete resume content.",
                    "error",
                )
                return render_index_with_context(context)

            if not job_posting:
                flash("Please paste the job posting.", "error")
                return render_index_with_context(context)

            if len(job_posting) < 50:
                flash(
                    "Job posting looks too short. Please paste the full posting for better results.",
                    "error",
                )
                return render_index_with_context(context)

            existing_snapshot = build_result_snapshot(context)

            try:
                if action == "regenerate" and existing_snapshot:
                    if not snapshot_exists(existing_snapshot, context["previous_versions"]):
                        context["previous_versions"].insert(0, existing_snapshot)

                job_title, company_name, tailored_bullets, cover_letter = (
                    generate_application_materials(resume_text, job_posting)
                )

                context["job_title"] = job_title
                context["company_name"] = company_name
                context["tailored_bullets"] = tailored_bullets
                context["cover_letter"] = cover_letter

                db = get_db()
                db.execute(
                    """
                    INSERT INTO applications (
                        user_id, resume_text, job_posting, tailored_bullets, cover_letter
                    )
                    VALUES (?, ?, ?, ?, ?)
                    """,
                    (
                        session["user_id"],
                        resume_text,
                        job_posting,
                        tailored_bullets,
                        cover_letter,
                    ),
                )
                db.commit()

                if action == "regenerate":
                    flash("New version generated and previous result kept below.", "success")
                else:
                    flash("Application materials generated and saved successfully.", "success")

            except RuntimeError as exc:
                flash(str(exc), "error")
            except sqlite3.Error:
                app.logger.exception("Database save failed after generation.")
                flash("The AI response was created, but saving to the database failed.", "error")
            except Exception:
                app.logger.exception("Unexpected generation error.")
                flash("An unexpected error occurred during generation.", "error")

            return render_index_with_context(context)

        flash("Unknown action requested.", "error")
        return render_index_with_context(context)

    return render_index_with_context(context)


@app.route("/history")
@login_required
def history():
    """
    Application history page only.

    This route intentionally reads only from the applications table and renders
    only history.html.
    """
    init_db()

    try:
        applications = fetch_application_history_for_user(session["user_id"])
    except sqlite3.Error:
        app.logger.exception("Application history query failed.")
        flash("Could not load application history.", "error")
        applications = []

    return render_template("history.html", applications=applications)


@app.route("/delete_application", methods=["POST"])
@login_required
def delete_application():
    """
    Delete one saved application-history item.
    """
    init_db()

    app_id = request.form.get("application_id")

    if not app_id:
        flash("Invalid application.", "error")
        return redirect(url_for("history"))

    try:
        db = get_db()
        db.execute(
            "DELETE FROM applications WHERE id = ? AND user_id = ?",
            (app_id, session["user_id"]),
        )
        db.commit()
        flash("Application deleted.", "success")
    except Exception:
        app.logger.exception("Failed to delete application.")
        flash("Could not delete application.", "error")

    return redirect(url_for("history"))


@app.route("/interview_history")
@login_required
def interview_history():
    """
    Interview session history page only.

    This route intentionally reads only from interview tables and renders only
    interview_history.html.
    """
    init_db()

    try:
        interview_sessions = build_interview_history_sessions(session["user_id"])
    except sqlite3.Error:
        app.logger.exception("Interview history query failed.")
        flash("Could not load interview history.", "error")
        interview_sessions = []

    return render_template("interview_history.html", interview_sessions=interview_sessions)


@app.route("/delete_interview_session", methods=["POST"])
@login_required
def delete_interview_session():
    """
    Delete one interview session and its related attempt rows.
    """
    init_db()

    session_id_raw = request.form.get("session_id", "").strip()

    try:
        session_id = int(session_id_raw)
    except (TypeError, ValueError):
        flash("Invalid interview session.", "error")
        return redirect(url_for("interview_history"))

    interview_session = get_interview_session_for_user(session["user_id"], session_id)
    if not interview_session:
        flash("Interview session not found.", "error")
        return redirect(url_for("interview_history"))

    try:
        db = get_db()
        db.execute("DELETE FROM interview_responses WHERE session_id = ?", (session_id,))
        db.execute(
            "DELETE FROM interview_sessions WHERE id = ? AND user_id = ?",
            (session_id, session["user_id"]),
        )
        db.commit()
        flash("Interview session deleted.", "success")
    except sqlite3.Error:
        app.logger.exception("Failed to delete interview session.")
        flash("Could not delete interview session.", "error")

    return redirect(url_for("interview_history"))


@app.route("/interview", methods=["GET", "POST"])
@login_required
def interview():
    """
    Interview Coach page.

    This route now uses session draft persistence so the page keeps its state
    when the user leaves and later returns.

    Easy-to-break area:
    This route relies on hidden form state in interview.html to preserve:
    - resume_text
    - job_posting
    - all_questions
    - current_interview_session_id

    It also now saves the current Interview Coach UI state into the Flask
    session, similar to how the Home page draft already works.
    """
    init_db()

    # Start with any saved Interview Coach draft so navigation away/back keeps state.
    context = load_interview_draft()

    if request.method == "GET":
        requested_session_id = request.args.get("session_id", "").strip()

        # If the user clicked "Open" from Interview History, load that DB-backed session.
        if requested_session_id:
            try:
                session_id = int(requested_session_id)
            except (TypeError, ValueError):
                session_id = None

            interview_session = get_interview_session_for_user(session["user_id"], session_id)
            if interview_session:
                context["resume_text"] = sanitize_model_text(interview_session["resume_text"])
                context["job_posting"] = sanitize_model_text(interview_session["job_posting"])
                context["interview_questions"] = sanitize_model_text(interview_session["questions_text"])
                context["all_questions"] = context["interview_questions"]
                context["current_interview_session_id"] = str(interview_session["id"])

                # These are intentionally cleared when opening a saved session so
                # the page does not show stale answer/feedback panels from a different session.
                context["generated_answer"] = ""
                context["selected_question"] = ""
                context["feedback_output"] = ""
                context["feedback_question"] = ""
                context["user_answer_text"] = ""
            else:
                flash("Interview session not found.", "error")

        return render_interview_with_context(context)

    # POST flow starts by rebuilding the context from hidden form fields.
    context = {
        "resume_text": request.form.get("resume_text", "").strip(),
        "job_posting": request.form.get("job_posting", "").strip(),
        "interview_questions": "",
        "generated_answer": "",
        "selected_question": "",
        "feedback_output": "",
        "feedback_question": "",
        "user_answer_text": "",
        "all_questions": request.form.get("all_questions", "").strip(),
        "current_interview_session_id": request.form.get("current_interview_session_id", "").strip(),
    }

    # Preserve whichever question block is currently on the page.
    context["interview_questions"] = context["all_questions"]

    action = request.form.get("action", "").strip()

    if action == "generate_questions":
        if not context["resume_text"] or not context["job_posting"]:
            flash("Please paste both your resume and job posting.", "error")
            return render_interview_with_context(context)

        if len(context["resume_text"]) < 50:
            flash("Resume text looks too short. Please paste more complete resume content.", "error")
            return render_interview_with_context(context)

        if len(context["job_posting"]) < 50:
            flash("Job posting looks too short. Please paste the full posting for better results.", "error")
            return render_interview_with_context(context)

        try:
            interview_questions = generate_interview_questions(
                context["resume_text"],
                context["job_posting"],
            )

            context["interview_questions"] = interview_questions
            context["all_questions"] = interview_questions
            context["current_interview_session_id"] = str(
                create_interview_session(
                    session["user_id"],
                    context["resume_text"],
                    context["job_posting"],
                    interview_questions,
                )
            )

            # Clear any stale answer/feedback panels from the previous run.
            context["generated_answer"] = ""
            context["selected_question"] = ""
            context["feedback_output"] = ""
            context["feedback_question"] = ""
            context["user_answer_text"] = ""

            flash("Interview session created and questions saved.", "success")
        except RuntimeError as exc:
            flash(str(exc), "error")
        except Exception:
            app.logger.exception("Unexpected interview question generation error.")
            flash("An unexpected error occurred while generating interview questions.", "error")

        return render_interview_with_context(context)

    elif action == "generate_answer":
        selected_question = request.form.get("question", "").strip()
        context["selected_question"] = selected_question
        context["interview_questions"] = context["all_questions"]

        if not selected_question:
            flash("No question selected.", "error")
            return render_interview_with_context(context)

        if not context["resume_text"] or not context["job_posting"]:
            flash("Resume text or job posting is missing.", "error")
            return render_interview_with_context(context)

        try:
            context["current_interview_session_id"] = str(
                ensure_interview_session_id(
                    context["current_interview_session_id"],
                    session["user_id"],
                    context["resume_text"],
                    context["job_posting"],
                    context["all_questions"],
                )
            )

            context["generated_answer"] = generate_interview_answer(
                selected_question,
                context["resume_text"],
                context["job_posting"],
            )

            save_interview_response_record(
                int(context["current_interview_session_id"]),
                selected_question,
                generated_answer=context["generated_answer"],
            )
        except RuntimeError as exc:
            flash(str(exc), "error")
        except sqlite3.Error:
            app.logger.exception("Interview answer save failed.")
            flash(
                "The answer was generated, but saving it to interview history failed. Restart the app and try again.",
                "error",
            )
        except Exception:
            app.logger.exception("Unexpected interview answer generation error.")
            flash("An unexpected error occurred while generating the answer.", "error")

        return render_interview_with_context(context)

    elif action == "get_feedback":
        feedback_question = request.form.get("question", "").strip()
        user_answer_text = request.form.get("user_answer_text", "").strip()

        context["feedback_question"] = feedback_question
        context["user_answer_text"] = user_answer_text
        context["interview_questions"] = context["all_questions"]

        if not feedback_question:
            flash("No question selected for feedback.", "error")
            return render_interview_with_context(context)

        if not user_answer_text:
            flash("Please type your answer before requesting feedback.", "error")
            return render_interview_with_context(context)

        if not context["resume_text"] or not context["job_posting"]:
            flash("Resume text or job posting is missing.", "error")
            return render_interview_with_context(context)

        try:
            context["current_interview_session_id"] = str(
                ensure_interview_session_id(
                    context["current_interview_session_id"],
                    session["user_id"],
                    context["resume_text"],
                    context["job_posting"],
                    context["all_questions"],
                )
            )

            context["feedback_output"] = generate_interview_feedback(
                feedback_question,
                user_answer_text,
                context["resume_text"],
                context["job_posting"],
            )

            save_interview_response_record(
                int(context["current_interview_session_id"]),
                feedback_question,
                user_answer_text=user_answer_text,
                feedback_output=context["feedback_output"],
            )
        except RuntimeError as exc:
            flash(str(exc), "error")
        except sqlite3.Error:
            app.logger.exception("Interview feedback save failed.")
            flash(
                "The feedback was generated, but saving it to interview history failed. Restart the app and try again.",
                "error",
            )
        except Exception:
            app.logger.exception("Unexpected interview feedback generation error.")
            flash("An unexpected error occurred while generating feedback.", "error")

        return render_interview_with_context(context)

    flash("Unknown action requested.", "error")
    return render_interview_with_context(context)


if __name__ == "__main__":
    with app.app_context():
        init_db()

    app.run(
        host="0.0.0.0",
        port=int(os.getenv("PORT", 5000)),
        debug=FLASK_DEBUG
    )




